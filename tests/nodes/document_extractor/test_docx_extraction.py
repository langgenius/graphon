import io
from unittest.mock import MagicMock

import pytest
from docx import Document

from graphon.file.enums import FileTransferMethod, FileType
from graphon.nodes.document_extractor import node as document_extractor_node
from graphon.nodes.document_extractor.exc import FileDownloadError


def test_extract_text_from_docx_keeps_paragraph_and_table_order() -> None:
    document = Document()
    document.add_paragraph("Intro")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Color"
    table.rows[1].cells[1].text = "Blue"
    document.add_paragraph("Outro")

    buffer = io.BytesIO()
    document.save(buffer)

    extracted = document_extractor_node._extract_text_from_docx(buffer.getvalue())

    assert extracted == (
        "Intro\n| Name | Value |\n| --- | --- |\n| Color | Blue |\n\nOutro"
    )


def test_download_file_content_requires_remote_url() -> None:
    file = MagicMock()
    file.transfer_method = FileTransferMethod.REMOTE_URL
    file.remote_url = None
    file.type = FileType.DOCUMENT

    with pytest.raises(FileDownloadError, match="Missing URL for remote file"):
        document_extractor_node._download_file_content(MagicMock(), file)
